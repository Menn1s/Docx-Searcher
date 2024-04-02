from colorama import Fore
import os
import docx
import argparse
import re

def search_keywords_in_docx_insensitive(file_path, keywords, header):
    doc = docx.Document(file_path)
    #text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    for i, paragraph in enumerate(doc.paragraphs):
        for keyword in keywords:
            if keyword.lower() in paragraph.text.lower():
                # set starting offset to iterate upwards and find the heading containing the matched text
                k = 1
                while not doc.paragraphs[i-k].style.name.startswith('Heading'):
                    #print(".", end="")
                    k += 1
                if header:
                    if doc.paragraphs[i-k].style.name.startswith('Heading') and (header in doc.paragraphs[i-k].text):
                        print(Fore.CYAN + "- Found some matches in " + Fore.YELLOW + file_path + Fore.RESET)
                        print("\t1. Found under heading: " + doc.paragraphs[i-k].text)
                        print("\t2. Matched Paragraph: " + doc.paragraphs[i].text)
                else:
                    if doc.paragraphs[i - k].style.name.startswith('Heading'):
                        print(Fore.CYAN + "- Found some matches in " + Fore.YELLOW + file_path + Fore.RESET)
                        print("\t1. Found under heading: " + doc.paragraphs[i - k].text)
                        print("\t2. Matched Paragraph: " + doc.paragraphs[i].text)


def search_keywords_in_docx_regex(file_path, keywords):
    doc = docx.Document(file_path)
    for i, paragraph in enumerate(doc.paragraphs):
        for keyword in keywords:
            pattern = re.compile(keyword, re.IGNORECASE)
            if pattern.search(paragraph.text):
                print(Fore.Cyan + "- Found some matches in " + Fore.YELLOW + file_path + Fore.RESET)
                # set starting offset to iterate upwards and find the heading containing the matched text
                k = 1
                while not doc.paragraphs[i - k].style.name.startswith('Heading'):
                    k += 1
                if doc.paragraphs[i - k].style.name.startswith('Heading'):
                    print("\t1. Found under heading: " + doc.paragraphs[i - k].text)
                print("\t2. Matched Paragraph: " + doc.paragraphs[i].text)

def search_keywords_in_docx(file_path, keywords):
    doc = docx.Document(file_path)
    for i, paragraph in enumerate(doc.paragraphs):
        for keyword in keywords:
            if keyword in paragraph.text:
                print(Fore.Cyan + "- Found some matches in " + Fore.YELLOW + file_path + Fore.RESET)
                # set starting offset to iterate upwards and find the heading containing the matched text
                k = 1
                while not doc.paragraphs[i-k].style.name.startswith('Heading'):
                    #print(".", end="")
                    k += 1
                if doc.paragraphs[i-k].style.name.startswith('Heading'):
                    print("\t1. Found under heading: " + doc.paragraphs[i-k].text)
                print("\t2. Matched Paragraph: " + doc.paragraphs[i].text)

def process_directory(directory_path, keywords, case_insensitive, extend_regex, header):
    for root, _, files in os.walk(directory_path):
        for file_name in files:
            if file_name.endswith('.docx') and not file_name.startswith('~'):
                file_path = os.path.join(root, file_name)
                if (case_insensitive):
                    try:
                        search_keywords_in_docx_insensitive(file_path, keywords, header)
                    except Exception as e:
                        print("Check that the file is not corrupted or password protected: " + file_path)
                        print(e)
                        pass
                        #print(f"Found keywords in: {file_path}")
                elif (extend_regex):
                    try:
                        search_keywords_in_docx_regex(file_path, keywords)
                    except Exception as e:
                        print("Check that the file is not corrupted or password protected: " + file_path)
                        print(e)
                        pass
                        #print(f"Found keywords in: {file_path}")
                else:
                    try:
                        search_keywords_in_docx_insensitive(file_path, keywords)
                    except Exception as e:
                        print("Check that the file is not corrupted or password protected: " + file_path)
                        print(e)
                        pass
                        #print(f"Found keywords in: {file_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Search for keywords in DOCX files in a directory.")
    exclusive_group = parser.add_mutually_exclusive_group()
    exclusive_group.add_argument("-i", "--case_insensitive", action="store_true", help="Search case insensitive")
    exclusive_group.add_argument("-e", "--extend_regex", action="store_true", help="Search with regex string")
    parser.add_argument("-H", "--header", help="Search for matching headers only")
    parser.add_argument("directory_path", help="Path to the directory containing DOCX files.")
    parser.add_argument("keywords", nargs="+", help="Keywords to search for in the DOCX files.")
    #parser.add_argument("-i", "--case_insensitive", action="store_true", help="Search case insensitive")
    #parser.add_argument("-e", "--extend_regex", action="store_true", help="Search with regex string")
    args = parser.parse_args()

    process_directory(args.directory_path, args.keywords, args.case_insensitive, args.extend_regex, args.header)
